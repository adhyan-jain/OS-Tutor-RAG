"""Extract text and metadata from Word (.docx) files into Documents."""

from __future__ import annotations

from pathlib import Path

import docx

from src.schemas import Document

_HEADING_LEVELS = {
    "Heading 1": 1,
    "Heading 2": 2,
    "Heading 3": 3,
}


def _heading_level(paragraph) -> int | None:
    return _HEADING_LEVELS.get(paragraph.style.name)


def extract_docx(file_path: Path) -> list[Document]:
    """Extract a DOCX file into one Document per section.

    Sections are delimited by Heading 1/2/3 paragraphs. Each Document's
    metadata records the section's heading, heading level, and the full
    heading path (e.g. ["Processes", "Scheduling", "Round Robin"]) down to
    that section. Any text preceding the first heading is emitted as its
    own Document with an empty section path.

    Args:
        file_path: Path to the .docx file.

    Returns:
        A list of Documents, one per section, each with section metadata.
    """
    document = docx.Document(file_path)

    documents: list[Document] = []
    heading_stack: list[tuple[int, str]] = []
    current_heading: str | None = None
    current_level: int | None = None
    current_body: list[str] = []
    section_index = 0

    def flush_section() -> None:
        nonlocal section_index
        text = "\n".join(current_body).strip()
        if not text and current_heading is None:
            return
        documents.append(
            Document(
                doc_id=f"{file_path.stem}__sec{section_index}",
                source_path=str(file_path),
                text=text,
                metadata={
                    "source_type": "docx",
                    "heading": current_heading,
                    "heading_level": current_level,
                    "section_path": [h for _, h in heading_stack],
                },
            )
        )
        section_index += 1

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        level = _heading_level(paragraph)

        if level is not None:
            if not text:
                continue
            flush_section()
            current_body = []

            while heading_stack and heading_stack[-1][0] >= level:
                heading_stack.pop()
            heading_stack.append((level, text))

            current_heading = text
            current_level = level
        else:
            if text:
                current_body.append(text)

    flush_section()

    return documents
